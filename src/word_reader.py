"""
Модуль для чтения .docx файлов и извлечения текста.
Извлекает текст из параграфов и таблиц.
"""

import os
import logging
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """Извлекает весь текст из .docx файла, включая таблицы."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    doc = Document(file_path)
    full_text = []

    # Извлекаем параграфы
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)

    # Извлекаем таблицы
    for table_idx, table in enumerate(doc.tables):
        full_text.append(f"\n--- Таблица {table_idx + 1} ---")
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ")
                row_data.append(cell_text)
            full_text.append(" | ".join(row_data))

    result = "\n".join(full_text)
    logger.info(f"Извлечено {len(result)} символов из {os.path.basename(file_path)}")
    return result


def get_all_docx_files(directory: str) -> list:
    """Возвращает список всех .docx файлов в директории."""
    files = []
    for f in os.listdir(directory):
        if f.endswith(".docx") and not f.startswith("~"):
            files.append(os.path.join(directory, f))
    logger.info(f"Найдено {len(files)} .docx файлов в {directory}")
    return files

# Тестовая функция
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    print("=" * 50)
    print("ТЕСТ МОДУЛЯ WORD_READER")
    print("=" * 50)
    
    test_dir = os.path.join(os.path.dirname(__file__), "input_docs")
    files = get_all_docx_files(test_dir)
    
    if files:
        print(f"✅ Найдено файлов: {len(files)}")
        for f in files[:5]:
            print(f"  - {os.path.basename(f)}")
    else:
        print(f"❌ Нет .docx файлов в {test_dir}")
    
    print("=" * 50)
